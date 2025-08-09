from docx import Document

def improve_resume(file_path, output_path, extracted_data):
    doc = Document()
    doc.add_heading(extracted_data.get('name', 'Name Unknown'), level=0)

    doc.add_heading('Contact Info', level=1)
    doc.add_paragraph(f"Email: {extracted_data.get('email')}")
    doc.add_paragraph(f"Phone: {extracted_data.get('phone')}")

    doc.add_heading('Skills', level=1)
    doc.add_paragraph(", ".join(extracted_data.get('skills', [])))

    doc.add_heading('Experience', level=1)
    for exp in extracted_data.get('experience', []):
        doc.add_paragraph(exp)

    doc.add_heading('Education', level=1)
    for edu in extracted_data.get('education', []):
        doc.add_paragraph(edu)

    doc.save(output_path)
